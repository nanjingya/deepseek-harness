#!/usr/bin/env python3
"""南鲸 PDF DOCX exporters for visual-fidelity and editable output modes.

The editable exporter uses the compact-reference-guide design preset with an
explicit A4/formal-Chinese-standard override.  Page images are used only for a
cover or genuinely image-only pages; OCR body content remains native Word text.

Project author: nanjingya.
"""
from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image


def platform_fonts(platform_name: str) -> tuple[str, str]:
    return ("SimSun", "SimHei") if platform_name == "win32" else ("Songti SC", "Heiti SC")


BODY_FONT, HEADING_FONT = platform_fonts(sys.platform)
LATIN_FONT = "Times New Roman"
INK = RGBColor(28, 31, 35)
MUTED = RGBColor(100, 105, 112)


def set_style_font(style, latin: str, east_asia: str, size: float, bold: bool = False, color: RGBColor = INK) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), latin)


def set_fonts(doc: Document) -> None:
    """Resolve the editable Word design preset into explicit style tokens."""
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, LATIN_FONT, BODY_FONT, 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Mm(7.4)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (15, 18, 10),
        "Heading 2": (13, 14, 7),
        "Heading 3": (11.5, 10, 5),
        "Heading 4": (10.5, 8, 4),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = styles[name]
        set_style_font(style, LATIN_FONT, HEADING_FONT, size, True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Mm(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, LATIN_FONT, BODY_FONT, 10.5)
        style.paragraph_format.left_indent = Mm(9.5)
        style.paragraph_format.first_line_indent = Mm(-4.7)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.widow_control = True

    caption = styles["Caption"]
    set_style_font(caption, LATIN_FONT, BODY_FONT, 9.5, False, INK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Mm(0)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_together = True


def force_run_fonts(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            style_name = paragraph.style.name if paragraph.style else ""
            east_asia = HEADING_FONT if style_name.startswith("Heading") else BODY_FONT
            run.font.name = LATIN_FONT
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            fonts.set(qn("w:ascii"), LATIN_FONT)
            fonts.set(qn("w:hAnsi"), LATIN_FONT)
            fonts.set(qn("w:eastAsia"), east_asia)
            fonts.set(qn("w:cs"), LATIN_FONT)


def set_a4_section(section, *, compact_cover: bool = False) -> None:
    section.page_width, section.page_height = Mm(210), Mm(297)
    if compact_cover:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Mm(5)
        section.header_distance = section.footer_distance = Mm(0)
    else:
        section.top_margin, section.bottom_margin = Mm(20), Mm(18)
        section.left_margin = section.right_margin = Mm(22)
        section.header_distance, section.footer_distance = Mm(9), Mm(9)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    current = OxmlElement("w:t"); current.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, current, end])
    paragraph.add_run(" 页")


def configure_running_furniture(section, title: str) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.text = title[:70]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    for run in hp.runs:
        run.font.name = LATIN_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), BODY_FONT)
    fp = section.footer.paragraphs[0]
    add_page_field(fp)
    for run in fp.runs:
        run.font.name = LATIN_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), BODY_FONT)
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), "1")


def parse_markdown_pages(markdown: Path) -> tuple[dict[str, str], dict[int, list[str]]]:
    metadata: dict[str, str] = {}
    pages: dict[int, list[str]] = {}
    current: int | None = None
    in_frontmatter = False
    for raw in markdown.read_text(encoding="utf-8").splitlines():
        stripped = raw.strip()
        if stripped == "---":
            in_frontmatter = not in_frontmatter
            continue
        if in_frontmatter:
            if ":" in stripped:
                key, value = stripped.split(":", 1)
                metadata[key.strip()] = value.strip().strip('"')
            continue
        marker = re.match(r"<!--\s*source_page:\s*(\d+)\s*-->", stripped)
        if not marker:
            marker = re.match(r">\s*\*\*PDF\s*原始页码：\*\*\s*(\d+)\s*$", stripped, re.I)
        if marker:
            current = int(marker.group(1))
            pages.setdefault(current, [])
            continue
        if current is not None:
            pages[current].append(raw)
    return metadata, pages


def add_image(doc: Document, path: Path, alt: str, *, max_width_mm: float = 166, max_height_mm: float = 222) -> None:
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = width_px / max(1, height_px)
    width_mm = min(max_width_mm, max_height_mm * ratio)
    height_mm = width_mm / ratio
    if height_mm > max_height_mm:
        height_mm = max_height_mm
        width_mm = height_mm * ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Mm(width_mm), height=Mm(height_mm))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)


COVER_SKIP_RE = re.compile(r"^(?:GM|刮涂层|查真伪)$")
STANDARD_NUMBER_RE = re.compile(
    r"^(?:GM|GB)\s*/?\s*[TZ]\s*\d+(?:\.\d+)?\s*(?:[—–-]\s*)?\d{4}$", re.I
)
TOC_ENTRY_RE = re.compile(r"^(.*?)\s*(?:……|\u2026{2}|\.{2,})\s*(\S+)\s*$")
SOURCE_PAGE_MARKER_RE = re.compile(
    r"(?:PDF\s*原始页码|source_page\s*:)",
    re.I,
)


def set_run_fonts(run, *, east_asia: str, size: float, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = INK
    run.font.name = LATIN_FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), LATIN_FONT)
    fonts.set(qn("w:hAnsi"), LATIN_FONT)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), LATIN_FONT)


def add_styled_paragraph(
    doc: Document,
    text: str,
    *,
    align,
    size: float,
    bold: bool = False,
    space_before: float = 0,
    space_after: float = 6,
    east_asia: str | None = None,
):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_fonts(run, east_asia=east_asia or HEADING_FONT, size=size, bold=bold)
    return p


def add_bottom_rule(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1C1F23")
    borders.append(bottom)
    p_pr.append(borders)


def hide_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tbl_pr.append(borders)


def normalize_cover_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = text.replace("a.gorithm", "algorithm")
    text = re.sub(r"\bcryptograph\b", "cryptography", text)
    text = re.sub(r"\b((?:GM|GB)/[TZ]\s*\d+(?:\.\d+)*)\s*-\s*(\d{4})\b", r"\1—\2", text, flags=re.I)
    text = re.sub(r"(备案号[：:])\s*(\d+)\s+[—–-]?\s*(\d{4})", r"\1\2—\3", text)
    text = re.sub(r"(?<=\d)(?=(?:实施|发布))", " ", text)
    return text


def load_page1_blocks(root: Path) -> list[dict]:
    json_path = root / "document.json"
    if not json_path.exists():
        return []
    try:
        payload = json.loads(json_path.read_text(encoding="utf-8"))
        page = next((item for item in payload.get("pages", []) if item.get("page") == 1), {})
        return list(page.get("blocks") or [])
    except (OSError, ValueError, TypeError):
        return []


def add_editable_cover(doc: Document, root: Path, title: str) -> bool:
    """Rebuild a standards cover from page-1 OCR instead of embedding the scan."""
    items: list[tuple[str, list[int]]] = []
    for block in load_page1_blocks(root):
        text = normalize_cover_text(block.get("text") or "")
        if not text or block.get("kind") in {"header_footer", "page_number"}:
            continue
        if COVER_SKIP_RE.match(text) or "刮涂层" in text or "查真伪" in text:
            continue
        bbox = block.get("bbox") or [0, 0, 0, 0]
        items.append((text, bbox))
    cjk = [text for text, _ in items if re.search(r"[一-鿿]", text)]
    if len(cjk) < 3:
        return False

    meta: list[str] = []
    series = ""
    number = ""
    english: list[str] = []
    dates: list[tuple[str, list[int]]] = []
    publisher: list[str] = []
    for text, bbox in items:
        compact = re.sub(r"\s+", "", text)
        if re.match(r"^ICS\b", text, re.I) or re.match(r"^[A-Z]\s*\d+$", text) or text.startswith("备案号"):
            meta.append(text)
        elif re.search(r"中华人民共和国.*(?:标准|技术文件)", text):
            series = text
        elif STANDARD_NUMBER_RE.match(compact):
            number = text
        elif re.search(r"\d{4}-\d{2}-\d{2}", text) and re.search(r"发布|实施", text):
            dates.append((text, bbox))
        elif "国家密码管理局" in text or compact == "发布":
            publisher.append(text)
        elif re.search(r"[A-Za-z]{4,}", text) and not re.search(r"[一-鿿]{4,}", text):
            english.append(text)

    set_a4_section(doc.sections[0])
    if doc.paragraphs:
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

    for line in meta:
        add_styled_paragraph(
            doc, line, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, space_after=2, east_asia=BODY_FONT,
        )
    if series:
        rule = add_styled_paragraph(
            doc, series, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True,
            space_before=18, space_after=4,
        )
        add_bottom_rule(rule)
    if number:
        add_styled_paragraph(
            doc, number, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, bold=True, space_before=6, space_after=28,
        )
    if title:
        add_styled_paragraph(
            doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, space_before=36, space_after=10,
        )
    if english:
        add_styled_paragraph(
            doc,
            " ".join(english),
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=11,
            space_after=36,
            east_asia=BODY_FONT,
        )
    if dates:
        ordered = sorted(dates, key=lambda item: item[1][0])
        table = doc.add_table(rows=1, cols=2)
        hide_table_borders(table)
        left_text = next((text for text, _ in ordered if "发布" in text), ordered[0][0])
        right_text = next((text for text, _ in ordered if "实施" in text), ordered[-1][0])
        left, right = table.rows[0].cells
        for cell, value, align in (
            (left, left_text, WD_ALIGN_PARAGRAPH.LEFT),
            (right, right_text, WD_ALIGN_PARAGRAPH.RIGHT),
        ):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = align
            paragraph.paragraph_format.first_line_indent = Mm(0)
            run = paragraph.add_run(value)
            set_run_fonts(run, east_asia=BODY_FONT, size=10.5)
    if publisher:
        publisher_text = re.sub(
            r"(国家密码管理局)\s*发布",
            r"\1  发布",
            "  ".join(dict.fromkeys(publisher)),
        )
        if "发布" not in publisher_text:
            publisher_text = f"{publisher_text}  发布"
        add_styled_paragraph(
            doc, publisher_text, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True,
            space_before=28, space_after=0,
        )
    return True


def add_cover(doc: Document, root: Path) -> bool:
    candidates = sorted(
        (root / "debug" / "pages").glob("page-*.jpg"),
        key=lambda path: int(path.stem.split("-")[-1]),
    )
    cover = candidates[0] if candidates and int(candidates[0].stem.split("-")[-1]) == 1 else root / "missing-cover.jpg"
    if not cover.exists():
        return False
    set_a4_section(doc.sections[0], compact_cover=True)
    p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    shape = p.add_run().add_picture(str(cover), width=Mm(200), height=Mm(282.8))
    shape._inline.docPr.set("descr", "原 PDF 封面")
    shape._inline.docPr.set("title", "原 PDF 封面")
    return True


def first_page_is_cover(root: Path, semantic_pages: dict[int, list[dict]]) -> bool:
    """Distinguish a real cover from a PDF whose first page is body content."""
    semantic = semantic_pages.get(1, [])
    if any(block.get("kind") == "full_page_image" for block in semantic):
        return True
    if any(block.get("kind") == "heading" for block in semantic):
        return False
    json_path = root / "document.json"
    if not json_path.exists():
        return False
    try:
        payload = json.loads(json_path.read_text(encoding="utf-8"))
        page = next((item for item in payload.get("pages", []) if item.get("page") == 1), {})
        raw = page.get("blocks", [])
    except (OSError, ValueError, TypeError):
        return False
    if not raw:
        return False
    heights = sorted(max(1, block["bbox"][3] - block["bbox"][1]) for block in raw if block.get("bbox"))
    if not heights:
        return False
    median_height = heights[len(heights) // 2]
    max_y = max(block["bbox"][3] for block in raw if block.get("bbox"))
    large_central_title = any(
        re.search(r"[一-鿿]", block.get("text", ""))
        and block["bbox"][3] - block["bbox"][1] >= median_height * 1.55
        and max_y * 0.20 <= (block["bbox"][1] + block["bbox"][3]) / 2 <= max_y * 0.75
        for block in raw if block.get("bbox")
    )
    cover_markers = sum(
        bool(re.search(r"(?:ICS\s*\d|CCS\s*[A-Z]|中华人民共和国.*(?:标准|技术文件)|发布|实施)", block.get("text", ""), re.I))
        for block in raw
    )
    return large_central_title or cover_markers >= 2


def meaningful_page(lines: list[str]) -> bool:
    return any(line.strip() and line.strip() != "---" for line in lines)


def page_is_full_image(lines: list[str]) -> bool:
    return any(re.match(r"!\[第\d+页整页图片\]", line.strip()) for line in lines)


def page_starts_major_section(lines: list[str]) -> bool:
    for line in lines:
        text = line.strip()
        if not text:
            continue
        heading = re.match(r"^#{1,2}\s+(.*)", text)
        return bool(heading and re.match(r"^(前言|引言|目录|附录|参考文献)", heading.group(1).replace(" ", "")))
    return False


def infer_running_title(root: Path, fallback: str) -> str:
    json_path = root / "document.json"
    if json_path.exists():
        try:
            payload = json.loads(json_path.read_text(encoding="utf-8"))
            first_page = next((page for page in payload.get("pages", []) if page.get("page") == 1), {})
            for block in first_page.get("blocks", []):
                text = block.get("text", "")
                if re.search(r"[A-Z]{1,4}\s*/\s*[A-Z]\s*\d{3,}", text, re.I):
                    return re.sub(r"\s*/\s*", "/", text).strip()
        except (OSError, ValueError, TypeError):
            pass
    return fallback


def load_semantic_pages(root: Path) -> dict[int, list[dict]]:
    """Load the shared layout model; older conversion folders use MD fallback."""
    json_path = root / "document.json"
    if not json_path.exists():
        return {}
    try:
        payload = json.loads(json_path.read_text(encoding="utf-8"))
        return {
            int(page["page"]): page.get("semantic_blocks", [])
            for page in payload.get("pages", [])
            if page.get("semantic_blocks") is not None
        }
    except (OSError, ValueError, TypeError, KeyError):
        return {}


def add_semantic_page(
    doc: Document, root: Path, blocks: list[dict], *, toc_heading_added: bool = False,
) -> bool:
    for block in blocks:
        kind = block.get("kind")
        text = (block.get("text") or "").strip()
        if kind == "heading" and text:
            markdown_level = int(block.get("level") or 2)
            level = 1 if markdown_level <= 2 else min(4, markdown_level - 1)
            p = doc.add_heading(text, level=level)
            if level == 1 and re.match(r"^(前言|引言|目录|目次|附录|参考文献)", text.replace(" ", "")):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "paragraph" and text:
            if SOURCE_PAGE_MARKER_RE.search(text):
                continue
            doc.add_paragraph(text, style="Normal")
        elif kind == "note" and text:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.first_line_indent = Mm(0)
            match = re.match(r"^(注\s*[：:])\s*(.*)", text)
            if match:
                p.add_run(match.group(1)).bold = True
                p.add_run(match.group(2))
            else:
                p.add_run(text)
        elif kind == "list_item" and text:
            numbered = re.match(r"^\d+[)）]", block.get("marker") or "")
            doc.add_paragraph(text, style="List Number" if numbered else "List Bullet")
        elif kind == "toc_entry" and text:
            if not toc_heading_added:
                heading = doc.add_heading("目 次", level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                toc_heading_added = True
            parsed = TOC_ENTRY_RE.match(text)
            label, page_number = (parsed.group(1), parsed.group(2)) if parsed else (text, "")
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Mm(0)
            p.paragraph_format.left_indent = Mm(7 * max(0, int(block.get("level") or 1) - 1))
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.tab_stops.add_tab_stop(Mm(166), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            run = p.add_run(label)
            set_run_fonts(run, east_asia=BODY_FONT, size=10.5)
            if page_number:
                page_run = p.add_run("\t" + page_number)
                set_run_fonts(page_run, east_asia=BODY_FONT, size=10.5)
        elif kind in {"figure", "table", "full_page_image"}:
            asset = block.get("asset")
            caption = block.get("caption") or "原文图片"
            if not asset:
                continue
            path = root / asset
            if not path.exists():
                continue
            add_image(doc, path, caption, max_height_mm=235 if kind == "full_page_image" else 212)
            if kind != "full_page_image":
                doc.add_paragraph(caption, style="Caption")
    return toc_heading_added


def semantic_page_starts_major_section(blocks: list[dict]) -> bool:
    first = next((block for block in blocks if block.get("kind") not in {None, "toc_entry"}), None)
    if not first or first.get("kind") != "heading":
        return False
    return bool(re.match(r"^(前言|引言|目录|附录|参考文献)", (first.get("text") or "").replace(" ", "")))


def add_markdown_page(doc: Document, root: Path, lines: list[str]) -> None:
    pending_image_alt: str | None = None
    in_details = False
    for raw in lines:
        line = raw.strip()
        if not line or line == "---":
            continue
        if SOURCE_PAGE_MARKER_RE.search(line):
            continue
        if line == "<details>":
            in_details = True
            continue
        if line == "</details>":
            in_details = False
            continue
        if in_details or line.startswith("<summary>") or re.fullmatch(r'<a id="source-page-\d+"></a>', line):
            continue
        image = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image:
            path = root / image.group(2)
            if path.exists():
                full_page = bool(re.match(r"第\d+页整页图片", image.group(1)))
                add_image(doc, path, image.group(1), max_height_mm=235 if full_page else 212)
                pending_image_alt = image.group(1)
            continue
        caption = re.fullmatch(r"\*\*(.*?)\*\*", line)
        if not caption:
            caption = re.fullmatch(r"\*(.*?)\*", line)
        if caption:
            text = caption.group(1).strip()
            if pending_image_alt and text == pending_image_alt:
                if not re.match(r"第\d+页整页图片", text):
                    doc.add_paragraph(text, style="Caption")
                pending_image_alt = None
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Mm(0)
                p.add_run(text).bold = True
            continue
        pending_image_alt = None
        heading = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading:
            markdown_level = len(heading.group(1))
            level = 1 if markdown_level <= 2 else min(4, markdown_level - 1)
            text = heading.group(2).strip()
            p = doc.add_heading(text, level=level)
            if level == 1 and re.match(r"^(前言|引言|目录|附录|参考文献)", text.replace(" ", "")):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        bullet = re.match(r"^-\s+(.*)", line)
        if bullet:
            text = re.sub(r"^[—-]\s*", "", bullet.group(1).strip())
            doc.add_paragraph(text, style="List Bullet")
            continue
        numbered = re.match(r"^\d+\.\s+(.*)", line)
        if numbered:
            doc.add_paragraph(numbered.group(1).strip(), style="List Number")
            continue
        note = re.match(r"^>\s*\*\*(注：)\*\*\s*(.*)", line)
        if note:
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.first_line_indent = Mm(0)
            p.add_run(note.group(1)).bold = True
            p.add_run(note.group(2))
            continue
        doc.add_paragraph(line, style="Normal")


def export_fidelity(pages: list[Path], output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Mm(5)
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1
    first = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    for index, image in enumerate(pages):
        paragraph = first if index == 0 else doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.add_run().add_picture(str(image), width=Mm(200), height=Mm(282.8))
        if index < len(pages) - 1:
            paragraph.add_run().add_break(WD_BREAK.PAGE)
    doc.core_properties.title = output.stem
    force_run_fonts(doc)
    doc.save(output)


def export_editable(markdown: Path, output: Path) -> None:
    root = markdown.parent
    doc = Document()
    set_fonts(doc)
    metadata, pages = parse_markdown_pages(markdown)
    semantic_pages = load_semantic_pages(root)
    title = metadata.get("title") or output.stem
    running_title = infer_running_title(root, title)

    has_cover = False
    if first_page_is_cover(root, semantic_pages):
        has_cover = add_editable_cover(doc, root, title) or add_cover(doc, root)
    if has_cover:
        body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    else:
        body_section = doc.sections[0]
        if doc.paragraphs:
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)
    set_a4_section(body_section)
    configure_running_furniture(body_section, running_title)

    body_started = False
    previous_was_full_image = False
    toc_heading_added = False
    page_numbers = sorted(semantic_pages) if semantic_pages else sorted(pages)
    for page_number in page_numbers:
        if page_number == 1 and has_cover:
            continue
        if semantic_pages:
            blocks = semantic_pages[page_number]
            if not blocks:
                continue
            full_image = any(block.get("kind") == "full_page_image" for block in blocks)
            major = semantic_page_starts_major_section(blocks)
            if body_started and (previous_was_full_image or full_image or major):
                doc.add_page_break()
            toc_heading_added = add_semantic_page(
                doc, root, blocks, toc_heading_added=toc_heading_added,
            )
        else:
            lines = pages[page_number]
            if not meaningful_page(lines):
                continue
            full_image = page_is_full_image(lines)
            if body_started and (previous_was_full_image or full_image or page_starts_major_section(lines)):
                doc.add_page_break()
            add_markdown_page(doc, root, lines)
        body_started = True
        previous_was_full_image = full_image

    doc.core_properties.title = title
    doc.core_properties.subject = "PDF OCR 可编辑版"
    doc.core_properties.comments = "正文为可编辑 OCR 文字；架构图和纯图页保留为图片。"
    doc.save(output)


def export(result_dir: Path, output: Path, mode: str) -> None:
    if mode == "fidelity":
        pages = sorted((result_dir / "debug" / "pages").glob("page-*.jpg"), key=lambda p: int(p.stem.split("-")[-1]))
        export_fidelity(pages, output)
    else:
        export_editable(result_dir / "document.md", output)
