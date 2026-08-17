#!/usr/bin/env python3
"""Cover reconstruction and TOC pairing checks for scanned Chinese standards."""

from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx_exporter import add_editable_cover, add_markdown_page, add_semantic_page, export_editable
from pdf_wiki_parser import (
    Line,
    build_semantic_blocks,
    build_toc_blocks,
    is_cover_noise,
    normalize_toc_label,
    toc_indent_level,
)


def toc(text: str, bbox: list[int]) -> Line:
    return Line(text, bbox, 50, 3, "toc")


class CoverAndTocTests(unittest.TestCase):
    def test_cover_noise_drops_logo_and_stamp(self) -> None:
        self.assertTrue(is_cover_noise("GM"))
        self.assertTrue(is_cover_noise("刮涂层 查真伪"))
        self.assertFalse(is_cover_noise("GM/T 0034-2014"))

    def test_cover_meta_and_dates_stay_separate(self) -> None:
        lines = [
            Line("ICS 35.040", [186, 56, 333, 89], 50, 1, "text"),
            Line("L 80", [186, 87, 254, 118], 30, 1, "text"),
            Line("备案号：44635 2014", [190, 115, 455, 149], 50, 1, "text"),
            Line("中华人民共和国密码行业标准", [220, 294, 1504, 383], 50, 1, "text"),
            Line("2014-02-13实施", [1270, 1941, 1545, 1982], 50, 1, "text"),
            Line("2014-02-13发布", [200, 1951, 472, 1989], 50, 1, "text"),
        ]
        texts = [block.text for block in build_semantic_blocks(lines, [], [], 1) if block.kind == "paragraph"]
        self.assertTrue(any(text.startswith("ICS") and "备案号" not in text for text in texts))
        self.assertTrue(any(text.startswith("L 80") for text in texts))
        self.assertTrue(any(text.startswith("备案号") for text in texts))
        self.assertTrue(any("发布" in text and "实施" not in text for text in texts))
        self.assertTrue(any("实施" in text and "发布" not in text for text in texts))

    def test_toc_pairs_unique_page_numbers_and_restores_hierarchy(self) -> None:
        lines = [
            toc("目", [753, 336, 791, 377]),
            toc("次", [883, 333, 927, 377]),
            toc("前言⋯", [193, 472, 295, 503]),
            toc("川", [1453, 469, 1473, 492]),
            toc("1 范围", [193, 523, 295, 554]),
            toc("2 规范性引用文件", [193, 574, 444, 605]),
            toc("1", [1463, 571, 1477, 594]),
            toc("术语和定义", [234, 621, 383, 656]),
            toc("4 缩咯语", [193, 676, 326, 707]),
            toc("5 证书认证系统", [196, 723, 417, 758]),
            toc("3", [1463, 724, 1480, 747]),
            toc("5.1 概述", [224, 778, 353, 809]),
            toc("3", [1463, 771, 1480, 795]),
            toc("6.4 KMC 与CA 的安全通信协议", [224, 1200, 620, 1235]),
            toc("14", [1463, 1204, 1480, 1228]),
            toc("密码算法、密码设备及接口", [193, 1250, 620, 1285]),
            toc("14", [1463, 1254, 1480, 1278]),
            toc("7.1 密码算法：", [224, 1300, 430, 1335]),
            toc("14", [1463, 1304, 1480, 1328]),
            toc("9.7人事管理制度", [230, 2053, 478, 2087]),
            toc("23", [1453, 2053, 1490, 2077]),
            toc("I", [1436, 2145, 1453, 2169]),
        ]
        blocks = build_toc_blocks(lines, 3)
        by_label = {block.text.split(" …… ")[0]: block for block in blocks}
        self.assertEqual(by_label["前言"].text, "前言 …… III")
        self.assertEqual(by_label["1 范围"].text, "1 范围 …… 1")
        self.assertEqual(by_label["2 规范性引用文件"].text, "2 规范性引用文件 …… 1")
        self.assertEqual(by_label["3 术语和定义"].text.split(" …… ")[0], "3 术语和定义")
        self.assertIn("4 缩略语", by_label)
        self.assertEqual(by_label["7 密码算法、密码设备及接口"].text, "7 密码算法、密码设备及接口 …… 14")
        self.assertEqual(by_label["9.7 人事管理制度"].text, "9.7 人事管理制度 …… 23")
        self.assertEqual(by_label["5.1 概述"].level, 2)
        self.assertEqual(by_label["前言"].level, 1)
        self.assertFalse(any(block.text.startswith("I") or block.text.endswith(" I") for block in blocks))

    def test_toc_label_normalization(self) -> None:
        self.assertEqual(normalize_toc_label("9.7人事管理制度"), "9.7 人事管理制度")
        self.assertEqual(normalize_toc_label("4 缩咯语"), "4 缩略语")
        self.assertEqual(toc_indent_level("A.1 当 RA 采用 C/S 模式时"), 2)

    def test_editable_cover_uses_native_text(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            (root / "document.json").write_text(json.dumps({
                "pages": [{
                    "page": 1,
                    "blocks": [
                        {"text": "ICS 35.040", "bbox": [186, 56, 333, 89], "kind": "text"},
                        {"text": "GM", "bbox": [1110, 70, 1440, 241], "kind": "header_footer"},
                        {"text": "L 80", "bbox": [186, 87, 254, 118], "kind": "text"},
                        {"text": "备案号：44635 2014", "bbox": [190, 115, 455, 149], "kind": "text"},
                        {"text": "中华人民共和国密码行业标准", "bbox": [220, 294, 1504, 383], "kind": "text"},
                        {"text": "GM/T 0034-2014", "bbox": [1171, 451, 1494, 489], "kind": "text"},
                        {"text": "2014-02-13发布", "bbox": [200, 1951, 472, 1989], "kind": "text"},
                        {"text": "2014-02-13实施", "bbox": [1270, 1941, 1545, 1982], "kind": "text"},
                        {"text": "国家密码管理局", "bbox": [638, 2077, 940, 2125], "kind": "text"},
                        {"text": "发布", "bbox": [1018, 2080, 1117, 2125], "kind": "text"},
                        {"text": "刮涂层 查真伪", "bbox": [241, 2182, 332, 2203], "kind": "header_footer"},
                    ],
                }],
            }, ensure_ascii=False), encoding="utf-8")
            doc = Document()
            title = "基于 SM2 密码算法的证书认证系统密码及其相关安全技术规范"
            self.assertTrue(add_editable_cover(doc, root, title))
            texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            self.assertIn("ICS 35.040", texts)
            self.assertIn("中华人民共和国密码行业标准", texts)
            self.assertIn("GM/T 0034—2014", texts)
            self.assertIn(title, texts)
            self.assertTrue(any("国家密码管理局" in text for text in texts))
            self.assertFalse(any(text == "GM" or "刮涂层" in text for text in texts))
            self.assertEqual(len(doc.inline_shapes), 0)
            self.assertEqual(len(doc.tables), 1)
            self.assertIn("发布", doc.tables[0].rows[0].cells[0].text)
            self.assertIn("实施", doc.tables[0].rows[0].cells[1].text)

    def test_toc_heading_is_written_once(self) -> None:
        doc = Document()
        first = [{"kind": "toc_entry", "text": "1 范围 …… 1", "level": 1}]
        second = [{"kind": "toc_entry", "text": "10.1 人员管理要求 …… 23", "level": 2}]
        added = add_semantic_page(doc, Path("."), first)
        added = add_semantic_page(doc, Path("."), second, toc_heading_added=added)
        headings = [paragraph.text for paragraph in doc.paragraphs if paragraph.style and paragraph.style.name.startswith("Heading")]
        self.assertEqual(headings.count("目 次"), 1)
        self.assertTrue(added)

    def test_word_omits_pdf_source_page_markers(self) -> None:
        doc = Document()
        add_markdown_page(doc, Path("."), [
            "> **PDF 原始页码：** 4",
            "—— PDF 原始页码：4 ——",
            "10.1 人员管理要求",
        ])
        texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        self.assertEqual(texts, ["10.1 人员管理要求"])
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            (root / "document.md").write_text(
                '---\ntitle: "示例"\n---\n\n'
                "> **PDF 原始页码：** 4\n\n"
                "正文一段。\n",
                encoding="utf-8",
            )
            output = root / "document.docx"
            export_editable(root / "document.md", output)
            exported = Document(str(output))
            body = "\n".join(paragraph.text for paragraph in exported.paragraphs)
            self.assertNotIn("原始页码", body)
            self.assertIn("正文一段。", body)


if __name__ == "__main__":
    unittest.main()
